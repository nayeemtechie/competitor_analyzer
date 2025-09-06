# src/competitor/reports/__init__.py
"""
Report generation for competitor analysis
"""

import os
import json
import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Optional
from dataclasses import asdict

from ..models import CompetitorIntelligence, CompetitorProfile
from ..config import CompetitorConfig

logger = logging.getLogger(__name__)

class ReportGenerator:
    """Generates various report formats for competitor analysis"""
    
    def __init__(self, config: CompetitorConfig, llm_provider):
        self.config = config
        self.llm_provider = llm_provider
        self.output_dir = Path(config.output_config['output_dir'])
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
    async def generate_pdf_report(self, intelligence: CompetitorIntelligence) -> str:
        """Generate comprehensive PDF report"""
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.graphics.shapes import Drawing
            from reportlab.graphics.charts.barcharts import VerticalBarChart
            from reportlab.graphics.charts.piecharts import Pie
            
        except ImportError:
            logger.error("ReportLab not installed. Install with: pip install reportlab")
            return await self.generate_json_report(intelligence)
        
        # Create filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"competitor_analysis_{timestamp}.pdf"
        filepath = self.output_dir / filename
        
        # Create PDF document
        doc = SimpleDocTemplate(str(filepath), pagesize=A4)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            textColor=colors.HexColor('#2c3e50')
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            spaceBefore=20,
            spaceAfter=12,
            textColor=colors.HexColor('#34495e')
        )
        
        # Title page
        story.append(Paragraph("Competitive Intelligence Report", title_style))
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph(f"Analysis Date: {datetime.now().strftime('%B %d, %Y')}", styles['Normal']))
        story.append(Paragraph(f"Competitors Analyzed: {len(intelligence.profiles)}", styles['Normal']))
        story.append(PageBreak())
        
        # Executive Summary
        story.append(Paragraph("Executive Summary", heading_style))
        exec_summary = await self._generate_executive_summary(intelligence)
        story.append(Paragraph(exec_summary, styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        
        # Market Overview
        story.append(Paragraph("Market Overview", heading_style))
        market_overview = self._generate_market_overview(intelligence)
        story.append(Paragraph(market_overview, styles['Normal']))
        story.append(PageBreak())
        
        # Competitive Landscape
        story.append(Paragraph("Competitive Landscape", heading_style))
        
        # Threat level summary
        threat_data = self._analyze_threat_levels(intelligence)
        if threat_data:
            story.append(self._create_threat_level_chart(threat_data))
            story.append(Spacer(1, 0.3*inch))
        
        # Competitor profiles
        for profile in intelligence.profiles:
            story.extend(self._create_competitor_section(profile, styles))
            story.append(PageBreak())
        
        # Competitive Matrix
        story.append(Paragraph("Competitive Feature Matrix", heading_style))
        matrix_table = self._create_feature_matrix_table(intelligence)
        if matrix_table:
            story.append(matrix_table)
            story.append(PageBreak())
        
        # Recommendations
        story.append(Paragraph("Strategic Recommendations", heading_style))
        recommendations = await self._generate_recommendations(intelligence)
        story.append(Paragraph(recommendations, styles['Normal']))
        
        # Build PDF
        doc.build(story)
        
        logger.info(f"PDF report generated: {filepath}")
        return str(filepath)
    
    async def generate_docx_report(self, intelligence: CompetitorIntelligence) -> str:
        """Generate Word document report"""
        try:
            from docx import Document
            from docx.shared import Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            return await self.generate_json_report(intelligence)
        
        # Create filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"competitor_analysis_{timestamp}.docx"
        filepath = self.output_dir / filename
        
        # Create document
        doc = Document()
        
        # Title
        title = doc.add_heading('Competitive Intelligence Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_paragraph(f"Analysis Date: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph(f"Competitors Analyzed: {len(intelligence.profiles)}")
        doc.add_page_break()
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        exec_summary = await self._generate_executive_summary(intelligence)
        doc.add_paragraph(exec_summary)
        
        # Market Overview
        doc.add_heading('Market Overview', level=1)
        market_overview = self._generate_market_overview(intelligence)
        doc.add_paragraph(market_overview)
        
        # Competitor Profiles
        doc.add_heading('Competitor Profiles', level=1)
        
        for profile in intelligence.profiles:
            # Competitor heading
            doc.add_heading(profile.name, level=2)
            
            # Basic info
            doc.add_paragraph(f"Website: {profile.website}")
            doc.add_paragraph(f"Threat Level: {profile.threat_level.value.title()}")
            
            if profile.target_markets:
                doc.add_paragraph(f"Target Markets: {', '.join(profile.target_markets)}")
            
            # Key features
            if profile.key_features:
                doc.add_heading('Key Features', level=3)
                for feature in profile.key_features[:10]:
                    doc.add_paragraph(f"• {feature}", style='List Bullet')
            
            # Funding info
            if profile.funding_info and profile.funding_info.total_funding:
                doc.add_heading('Funding Information', level=3)
                doc.add_paragraph(f"Total Funding: {profile.funding_info.total_funding}")
                if profile.funding_info.last_round_amount:
                    doc.add_paragraph(f"Last Round: {profile.funding_info.last_round_amount}")
            
            doc.add_page_break()
        
        # Feature Matrix
        doc.add_heading('Competitive Feature Matrix', level=1)
        matrix_data = self._create_feature_matrix_data(intelligence)
        if matrix_data:
            table = doc.add_table(rows=len(matrix_data), cols=len(matrix_data[0]))
            table.style = 'Table Grid'
            
            for i, row in enumerate(matrix_data):
                for j, cell_value in enumerate(row):
                    table.cell(i, j).text = str(cell_value)
        
        # Recommendations
        doc.add_heading('Strategic Recommendations', level=1)
        recommendations = await self._generate_recommendations(intelligence)
        doc.add_paragraph(recommendations)
        
        # Save document
        doc.save(str(filepath))
        
        logger.info(f"DOCX report generated: {filepath}")
        return str(filepath)
    
    async def generate_json_report(self, intelligence: CompetitorIntelligence) -> str:
        """Generate JSON report with all data"""
        # Create filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"competitor_analysis_{timestamp}.json"
        filepath = self.output_dir / filename
        
        # Convert intelligence to dict
        report_data = {
            'metadata': {
                'generated_at': datetime.now().isoformat(),
                'total_competitors': len(intelligence.profiles),
                'analysis_version': '1.0',
                'config_summary': self._get_config_summary()
            },
            'executive_summary': await self._generate_executive_summary(intelligence),
            'market_overview': self._generate_market_overview(intelligence),
            'threat_analysis': self._analyze_threat_levels(intelligence),
            'competitors': [],
            'competitive_matrix': self._create_feature_matrix_data(intelligence),
            'recommendations': await self._generate_recommendations(intelligence)
        }
        
        # Add competitor data
        for profile in intelligence.profiles:
            competitor_data = asdict(profile)
            # Clean up data for JSON serialization
            competitor_data = self._clean_for_json(competitor_data)
            report_data['competitors'].append(competitor_data)
        
        # Save JSON
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(report_data, f, indent=2, default=str, ensure_ascii=False)
        
        logger.info(f"JSON report generated: {filepath}")
        return str(filepath)
    
    def _clean_for_json(self, data: Any) -> Any:
        """Clean data for JSON serialization"""
        if isinstance(data, dict):
            return {k: self._clean_for_json(v) for k, v in data.items()}
        elif isinstance(data, list):
            return [self._clean_for_json(item) for item in data]
        elif hasattr(data, '__dict__'):
            return self._clean_for_json(asdict(data))
        else:
            return data
    
    async def _generate_executive_summary(self, intelligence: CompetitorIntelligence) -> str:
        """Generate executive summary using LLM"""
        if not intelligence.profiles:
            return "No competitor data available for analysis."
        
        # Prepare data for LLM
        competitor_summaries = []
        for profile in intelligence.profiles:
            summary_data = {
                'name': profile.name,
                'threat_level': profile.threat_level.value,
                'key_features_count': len(profile.key_features) if profile.key_features else 0,
                'funding': profile.funding_info.total_funding if profile.funding_info else 'Unknown',
                'target_markets': profile.target_markets or []
            }
            competitor_summaries.append(summary_data)
        
        try:
            system_prompt = """You are a senior business analyst creating an executive summary of competitive landscape analysis. 
            Focus on strategic implications, key threats, and actionable insights for C-level executives."""
            
            user_prompt = f"""
            Create an executive summary based on analysis of {len(intelligence.profiles)} key competitors:
            
            Competitor Data:
            {json.dumps(competitor_summaries, indent=2)}
            
            Structure the summary with:
            1. Market Overview (2-3 sentences)
            2. Primary Competitive Threats (3 key points)
            3. Strategic Opportunities (2-3 points)
            4. Recommended Actions (3 specific recommendations)
            
            Keep it concise, executive-friendly, and actionable. Maximum 300 words.
            """
            
            return await self.llm_provider.achat(system_prompt, user_prompt, model="gpt-4o")
            
        except Exception as e:
            logger.warning(f"LLM summary generation failed: {e}")
            
            # Fallback summary
            high_threat_count = len([p for p in intelligence.profiles 
                                   if p.threat_level.value in ['high', 'critical']])
            
            return f"""
            COMPETITIVE LANDSCAPE EXECUTIVE SUMMARY
            
            Market Analysis: Analyzed {len(intelligence.profiles)} key competitors in the ecommerce search market.
            
            Key Findings:
            • {high_threat_count} competitors pose high competitive threat requiring immediate attention
            • Market shows continued innovation with AI/ML capabilities becoming standard
            • Strong funding activity indicates healthy market growth and competitive intensity
            
            Strategic Implications:
            • Accelerated product development needed to maintain competitive differentiation
            • Customer retention strategies critical as competition intensifies
            • Technology partnerships may provide competitive advantages
            
            Recommended Actions:
            • Conduct quarterly competitive feature analysis
            • Strengthen customer success and retention programs
            • Evaluate strategic partnerships or acquisition opportunities
            """
    
    def _generate_market_overview(self, intelligence: CompetitorIntelligence) -> str:
        """Generate market overview section"""
        if not intelligence.profiles:
            return "No market data available."
        
        # Market segment analysis
        all_markets = []
        for profile in intelligence.profiles:
            if profile.target_markets:
                all_markets.extend(profile.target_markets)
        
        market_segments = {}
        for market in all_markets:
            market_segments[market] = market_segments.get(market, 0) + 1
        
        # Funding analysis
        total_funding = 0
        funded_companies = 0
        for profile in intelligence.profiles:
            if profile.funding_info and profile.funding_info.total_funding:
                funded_companies += 1
                # Simple parsing for demonstration
                funding_str = profile.funding_info.total_funding
                if 'B' in funding_str:
                    total_funding += 1000  # Billions in millions
                elif 'M' in funding_str:
                    try:
                        amount = float(funding_str.replace('$', '').replace('M', '').replace(',', ''))
                        total_funding += amount
                    except:
                        pass
        
        overview = f"""
        The ecommerce search market shows strong competitive activity with {len(intelligence.profiles)} major players analyzed.
        
        Market Segmentation:
        • Primary focus on {', '.join(list(market_segments.keys())[:3])} segments
        • Enterprise market shows highest competitive intensity
        • Mid-market segment experiencing rapid growth
        
        Market Dynamics:
        • Total analyzed funding: ~${total_funding:.0f}M across {funded_companies} companies
        • Technology focus shifting toward AI/ML capabilities
        • Platform consolidation creating comprehensive solution providers
        
        Competitive Intensity: HIGH - Multiple well-funded players competing across overlapping segments
        """
        
        return overview.strip()
    
    def _analyze_threat_levels(self, intelligence: CompetitorIntelligence) -> Dict[str, Any]:
        """Analyze threat levels across competitors"""
        threat_counts = {'low': 0, 'medium': 0, 'high': 0, 'critical': 0}
        threat_details = {'high_threats': [], 'emerging_threats': []}
        
        for profile in intelligence.profiles:
            threat_level = profile.threat_level.value
            threat_counts[threat_level] += 1
            
            if threat_level in ['high', 'critical']:
                threat_details['high_threats'].append({
                    'name': profile.name,
                    'threat_level': threat_level,
                    'key_factors': self._get_threat_factors(profile)
                })
            elif threat_level == 'medium' and self._is_emerging_threat(profile):
                threat_details['emerging_threats'].append(profile.name)
        
        return {
            'distribution': threat_counts,
            'details': threat_details,
            'total_analyzed': len(intelligence.profiles)
        }
    
    def _get_threat_factors(self, profile: CompetitorProfile) -> List[str]:
        """Get key threat factors for a competitor"""
        factors = []
        
        if profile.funding_info and profile.funding_info.total_funding:
            if 'B' in profile.funding_info.total_funding or '500M' in profile.funding_info.total_funding:
                factors.append("Strong funding position")
        
        if profile.key_features and len(profile.key_features) > 15:
            factors.append("Comprehensive feature set")
        
        if profile.job_postings and len(profile.job_postings) > 10:
            factors.append("Aggressive hiring")
        
        return factors[:3]
    
    def _is_emerging_threat(self, profile: CompetitorProfile) -> bool:
        """Check if competitor is an emerging threat"""
        # Simple heuristics for emerging threats
        if profile.funding_info and profile.funding_info.last_round_date:
            try:
                from datetime import datetime, timedelta
                from dateutil import parser
                last_round = parser.parse(profile.funding_info.last_round_date)
                if (datetime.now() - last_round).days < 365:
                    return True
            except:
                pass
        
        return False
    
    def _create_feature_matrix_data(self, intelligence: CompetitorIntelligence) -> List[List[str]]:
        """Create feature comparison matrix data"""
        if not intelligence.profiles:
            return []
        
        # Key features to compare
        key_features = [
            'Search Relevance', 'Autocomplete', 'Faceted Search', 'Personalization',
            'A/B Testing', 'Analytics', 'API Access', 'AI/ML', 'Real-time'
        ]
        
        # Create matrix
        matrix = []
        
        # Header row
        header = ['Feature'] + [profile.name for profile in intelligence.profiles]
        matrix.append(header)
        
        # Feature rows
        for feature in key_features:
            row = [feature]
            for profile in intelligence.profiles:
                has_feature = self._check_feature_presence(feature, profile)
                row.append('✓' if has_feature else '✗')
            matrix.append(row)
        
        return matrix
    
    def _check_feature_presence(self, feature_name: str, profile: CompetitorProfile) -> bool:
        """Check if competitor has a specific feature"""
        if not profile.key_features:
            return False
        
        feature_keywords = {
            'Search Relevance': ['relevance', 'ranking', 'scoring'],
            'Autocomplete': ['autocomplete', 'suggestion', 'typeahead'],
            'Faceted Search': ['facet', 'filter', 'refinement'],
            'Personalization': ['personalization', 'recommendation', 'custom'],
            'A/B Testing': ['a/b', 'split test', 'experiment'],
            'Analytics': ['analytics', 'reporting', 'metrics'],
            'API Access': ['api', 'rest', 'graphql'],
            'AI/ML': ['ai', 'machine learning', 'ml', 'neural'],
            'Real-time': ['real-time', 'live', 'instant', 'streaming']
        }
        
        keywords = feature_keywords.get(feature_name, [feature_name.lower()])
        
        for feature in profile.key_features:
            feature_lower = feature.lower()
            if any(keyword in feature_lower for keyword in keywords):
                return True
        
        return False
    
    def _create_feature_matrix_table(self, intelligence: CompetitorIntelligence):
        """Create feature matrix table for PDF"""
        try:
            from reportlab.platypus import Table, TableStyle
            from reportlab.lib import colors
            
            matrix_data = self._create_feature_matrix_data(intelligence)
            if not matrix_data:
                return None
            
            table = Table(matrix_data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            return table
            
        except ImportError:
            return None
    
    def _create_threat_level_chart(self, threat_data: Dict[str, Any]):
        """Create threat level chart for PDF"""
        try:
            from reportlab.graphics.shapes import Drawing
            from reportlab.graphics.charts.piecharts import Pie
            from reportlab.lib import colors
            
            distribution = threat_data['distribution']
            
            # Create drawing
            drawing = Drawing(400, 200)
            
            # Create pie chart
            pie = Pie()
            pie.x = 50
            pie.y = 50
            pie.width = 100
            pie.height = 100
            
            # Data and labels
            data = []
            labels = []
            chart_colors = []
            
            color_map = {
                'low': colors.green,
                'medium': colors.yellow,
                'high': colors.orange,
                'critical': colors.red
            }
            
            for level, count in distribution.items():
                if count > 0:
                    data.append(count)
                    labels.append(f"{level.title()}: {count}")
                    chart_colors.append(color_map.get(level, colors.grey))
            
            pie.data = data
            pie.labels = labels
            pie.slices.strokeColor = colors.white
            pie.slices.strokeWidth = 1
            
            # Set colors
            for i, color in enumerate(chart_colors):
                pie.slices[i].fillColor = color
            
            drawing.add(pie)
            return drawing
            
        except ImportError:
            return None
    
    def _create_competitor_section(self, profile: CompetitorProfile, styles) -> List:
        """Create competitor profile section for PDF"""
        from reportlab.platypus import Paragraph, Spacer
        from reportlab.lib.units import inch
        
        section = []
        
        # Competitor name
        section.append(Paragraph(profile.name, styles['Heading2']))
        section.append(Spacer(1, 0.1*inch))
        
        # Basic info
        basic_info = f"""
        <b>Website:</b> {profile.website}<br/>
        <b>Threat Level:</b> {profile.threat_level.value.title()}<br/>
        """
        
        if profile.target_markets:
            basic_info += f"<b>Target Markets:</b> {', '.join(profile.target_markets)}<br/>"
        
        if profile.funding_info and profile.funding_info.total_funding:
            basic_info += f"<b>Total Funding:</b> {profile.funding_info.total_funding}<br/>"
        
        section.append(Paragraph(basic_info, styles['Normal']))
        section.append(Spacer(1, 0.2*inch))
        
        # Key features
        if profile.key_features:
            section.append(Paragraph("Key Features:", styles['Heading3']))
            features_text = "<br/>".join([f"• {feature}" for feature in profile.key_features[:8]])
            section.append(Paragraph(features_text, styles['Normal']))
            section.append(Spacer(1, 0.2*inch))
        
        # Recent activity
        activity_text = ""
        if profile.recent_news:
            activity_text += f"Recent News: {len(profile.recent_news)} mentions<br/>"
        if profile.job_postings:
            activity_text += f"Active Job Postings: {len(profile.job_postings)}<br/>"
        if profile.github_activity:
            activity_text += f"GitHub Repos: {profile.github_activity.public_repos}<br/>"
        
        if activity_text:
            section.append(Paragraph("Recent Activity:", styles['Heading3']))
            section.append(Paragraph(activity_text, styles['Normal']))
        
        return section
    
    async def _generate_recommendations(self, intelligence: CompetitorIntelligence) -> str:
        """Generate strategic recommendations using LLM"""
        if not intelligence.profiles:
            return "No data available for recommendations."
        
        try:
            # Prepare competitive analysis data
            analysis_summary = {
                'total_competitors': len(intelligence.profiles),
                'high_threats': len([p for p in intelligence.profiles 
                                   if p.threat_level.value in ['high', 'critical']]),
                'well_funded': len([p for p in intelligence.profiles 
                                  if p.funding_info and p.funding_info.total_funding and 
                                  ('B' in p.funding_info.total_funding or '100M' in p.funding_info.total_funding)]),
                'feature_leaders': [p.name for p in intelligence.profiles 
                                  if p.key_features and len(p.key_features) > 15][:3],
                'active_hirers': [p.name for p in intelligence.profiles 
                                if p.job_postings and len(p.job_postings) > 10][:3]
            }
            
            system_prompt = """You are a strategic business consultant providing actionable recommendations 
            based on competitive intelligence analysis. Focus on practical, prioritized actions."""
            
            user_prompt = f"""
            Based on competitive analysis of {analysis_summary['total_competitors']} competitors:
            
            Key Findings:
            - {analysis_summary['high_threats']} high-threat competitors
            - {analysis_summary['well_funded']} well-funded competitors
            - Feature leaders: {', '.join(analysis_summary['feature_leaders'])}
            - Active hiring: {', '.join(analysis_summary['active_hirers'])}
            
            Provide 5 strategic recommendations with:
            1. Specific action
            2. Timeline (immediate/3-month/6-month)
            3. Expected impact
            
            Format as numbered list, keep concise and actionable.
            """
            
            return await self.llm_provider.achat(system_prompt, user_prompt, model="gpt-4o")
            
        except Exception as e:
            logger.warning(f"LLM recommendations generation failed: {e}")
            
            # Fallback recommendations
            high_threat_count = len([p for p in intelligence.profiles 
                                   if p.threat_level.value in ['high', 'critical']])
            
            return f"""
            STRATEGIC RECOMMENDATIONS
            
            1. IMMEDIATE (Next 30 days)
               • Conduct detailed feature gap analysis against top {min(3, high_threat_count)} competitors
               • Review and strengthen customer retention programs
               • Assess pricing competitiveness
            
            2. SHORT-TERM (3 months)
               • Accelerate AI/ML feature development to match market standards
               • Expand sales team to compete for enterprise deals
               • Develop competitive battle cards for sales team
            
            3. MEDIUM-TERM (6 months)
               • Consider strategic partnerships to fill technology gaps
               • Evaluate acquisition opportunities for complementary capabilities
               • Strengthen developer ecosystem and API offerings
            
            4. LONG-TERM (12 months)
               • Assess market expansion opportunities in underserved verticals
               • Build thought leadership through content and conference presence
               • Develop platform ecosystem to increase switching costs
            
            5. CONTINUOUS
               • Implement quarterly competitive monitoring and analysis
               • Track competitor product releases and feature updates
               • Monitor funding rounds and team expansion activities
            """
    
    def _get_config_summary(self) -> Dict[str, Any]:
        """Get configuration summary for report metadata"""
        return {
            'analysis_depth': self.config.analysis_config.get('depth_level', 'standard'),
            'data_sources_enabled': [
                source for source, enabled in self.config.data_sources_config.items()
                if (isinstance(enabled, dict) and enabled.get('enabled')) or 
                   (isinstance(enabled, bool) and enabled)
            ],
            'output_format': self.config.output_config.get('formats', ['pdf']),
            'competitors_configured': len(self.config.competitors)
        }