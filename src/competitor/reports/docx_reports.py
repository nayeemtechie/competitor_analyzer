# src/competitor/reports/docx_reports.py
"""
Word document report generator for competitor analysis
"""

import os
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Optional
import logging

logger = logging.getLogger(__name__)

class DOCXReportGenerator:
    """Generates Word document reports"""
    
    def __init__(self, config, llm_provider):
        self.config = config
        self.llm_provider = llm_provider
        self.output_dir = Path(config.output_config.get('output_dir', 'competitor_reports'))
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # DOCX configuration
        self.docx_config = config.output_config.get('docx', {})
        self.template = self.docx_config.get('template')
        self.include_toc = self.docx_config.get('include_toc', True)
        self.include_charts = self.docx_config.get('include_charts', True)
    
    async def generate_report(self, intelligence) -> str:
        """Generate Word document report"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
            from docx.enum.style import WD_STYLE_TYPE
            from docx.oxml.shared import OxmlElement, qn
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            return await self._generate_simple_text_report(intelligence)
        
        # Create document
        if self.template and Path(self.template).exists():
            doc = Document(self.template)
        else:
            doc = Document()
        
        # Configure styles
        self._configure_document_styles(doc)
        
        # Title page
        self._create_title_page(doc, intelligence)
        doc.add_page_break()
        
        # Table of Contents
        if self.include_toc:
            self._create_table_of_contents(doc, intelligence)
            doc.add_page_break()
        
        # Executive Summary
        await self._create_executive_summary(doc, intelligence)
        doc.add_page_break()
        
        # Competitive Threat Matrix
        self._create_threat_matrix(doc, intelligence)
        doc.add_page_break()
        
        # Market Analysis
        await self._create_market_analysis(doc, intelligence)
        doc.add_page_break()
        
        # Individual Competitor Profiles
        for profile in intelligence.profiles:
            self._create_competitor_profile(doc, profile)
            doc.add_page_break()
        
        # Strategic Recommendations
        await self._create_recommendations(doc, intelligence)
        doc.add_page_break()
        
        # Appendix
        self._create_appendix(doc, intelligence)
        
        # Save document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"competitor_analysis_{timestamp}.docx"
        filepath = self.output_dir / filename
        doc.save(str(filepath))
        
        logger.info(f"DOCX report generated: {filepath}")
        return str(filepath)
    
    def _configure_document_styles(self, doc):
        """Configure document styles"""
        try:
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            styles = doc.styles
            
            # Title style
            if 'Report Title' not in [s.name for s in styles]:
                title_style = styles.add_style('Report Title', WD_STYLE_TYPE.PARAGRAPH)
                title_format = title_style.paragraph_format
                title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_format.space_after = Pt(30)
                
                title_font = title_style.font
                title_font.name = 'Calibri'
                title_font.size = Pt(24)
                title_font.bold = True
                title_font.color.rgb = RGBColor(52, 152, 219)
            
            # Heading styles
            heading1 = styles['Heading 1']
            heading1.font.color.rgb = RGBColor(52, 152, 219)
            heading1.font.size = Pt(18)
            
            heading2 = styles['Heading 2']
            heading2.font.color.rgb = RGBColor(52, 152, 219)
            heading2.font.size = Pt(14)
            
        except Exception as e:
            logger.warning(f"Could not configure styles: {e}")
    
    def _create_title_page(self, doc, intelligence):
        """Create title page"""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Main title
        title = doc.add_heading('Competitive Intelligence Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_paragraph('Ecommerce Search Market Analysis')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Analysis details
        doc.add_paragraph()  # Empty line
        doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
        doc.add_paragraph(f'Competitors Analyzed: {len(intelligence.profiles)}')
        
        if intelligence.config:
            doc.add_paragraph(f'Analysis Depth: {intelligence.config.depth_level.value.title()}')
        
        # Competitor list
        doc.add_paragraph()
        doc.add_paragraph('Companies Included:').bold = True
        for profile in intelligence.profiles:
            doc.add_paragraph(f'• {profile.name}', style='List Bullet')
        
        # Disclaimer
        doc.add_paragraph()
        disclaimer = doc.add_paragraph()
        disclaimer.add_run('This report contains confidential competitive intelligence. ')
        disclaimer.add_run('The information herein is intended for internal strategic planning purposes only. ')
        disclaimer.add_run('Distribution should be limited to authorized personnel.')
        disclaimer.italic = True
    
    def _create_table_of_contents(self, doc, intelligence):
        """Create table of contents"""
        doc.add_heading('Table of Contents', level=1)
        
        toc_items = [
            '1. Executive Summary',
            '2. Competitive Threat Matrix',
            '3. Market Landscape Analysis',
            '4. Individual Competitor Profiles'
        ]
        
        # Add each competitor to TOC
        for i, profile in enumerate(intelligence.profiles, 1):
            toc_items.append(f'   4.{i} {profile.name}')
        
        toc_items.extend([
            '5. Strategic Recommendations',
            '6. Appendix'
        ])
        
        for item in toc_items:
            doc.add_paragraph(item, style='List Number')
    
    async def _create_executive_summary(self, doc, intelligence):
        """Create executive summary section"""
        doc.add_heading('Executive Summary', level=1)
        
        # Generate executive summary
        executive_summary = await self._generate_executive_summary_content(intelligence)
        
        # Split into paragraphs and add to document
        paragraphs = executive_summary.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        # Key metrics summary
        doc.add_heading('Key Market Metrics', level=2)
        
        metrics = [
            f'Total competitors analyzed: {len(intelligence.profiles)}',
            f'High-threat competitors: {len([p for p in intelligence.profiles if p.threat_level and p.threat_level.value in ["high", "critical"]])}',
            f'Well-funded competitors: {len([p for p in intelligence.profiles if p.funding_info and p.funding_info.total_funding and "B" in p.funding_info.total_funding])}',
            f'Active hiring competitors: {len([p for p in intelligence.profiles if p.job_postings and len(p.job_postings) > 10])}'
        ]
        
        for metric in metrics:
            doc.add_paragraph(metric, style='List Bullet')
    
    def _create_threat_matrix(self, doc, intelligence):
        """Create threat matrix section"""
        from docx.shared import Inches
        
        doc.add_heading('Competitive Threat Matrix', level=1)
        
        # Create threat level table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        table.autofit = False
        
        # Set column widths
        for i, width in enumerate([Inches(1.5), Inches(1), Inches(1), Inches(1.8), Inches(1.8)]):
            table.columns[i].width = width
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Competitor'
        hdr_cells[1].text = 'Threat Level'
        hdr_cells[2].text = 'Funding'
        hdr_cells[3].text = 'Key Strengths'
        hdr_cells[4].text = 'Primary Concerns'
        
        # Make header bold
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Data rows
        for profile in intelligence.profiles:
            row_cells = table.add_row().cells
            row_cells[0].text = profile.name
            row_cells[1].text = profile.threat_level.value.title() if profile.threat_level else 'Medium'
            row_cells[2].text = self._get_funding_status(profile)
            row_cells[3].text = self._get_key_strengths(profile)
            row_cells[4].text = self._get_primary_concerns(profile)
        
        # Threat level distribution
        doc.add_heading('Threat Level Distribution', level=2)
        
        threat_counts = {'critical': 0, 'high': 0, 'medium': 0, 'low': 0}
        for profile in intelligence.profiles:
            level = profile.threat_level.value if profile.threat_level else 'medium'
            threat_counts[level] += 1
        
        for level, count in threat_counts.items():
            if count > 0:
                percentage = (count / len(intelligence.profiles)) * 100
                doc.add_paragraph(f'{level.title()}: {count} competitors ({percentage:.1f}%)', style='List Bullet')
    
    async def _create_market_analysis(self, doc, intelligence):
        """Create market analysis section"""
        doc.add_heading('Market Landscape Analysis', level=1)
        
        # Market overview
        doc.add_heading('Market Overview', level=2)
        market_analysis = await self._generate_market_analysis_content(intelligence)
        
        paragraphs = market_analysis.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        # Competitive dynamics
        doc.add_heading('Competitive Dynamics', level=2)
        
        dynamics = [
            f'Market consolidation: {len([p for p in intelligence.profiles if p.funding_info and p.funding_info.total_funding])} of {len(intelligence.profiles)} competitors are well-funded',
            f'Innovation focus: {len([p for p in intelligence.profiles if p.key_features and any("ai" in f.lower() for f in p.key_features)])} competitors emphasize AI/ML capabilities',
            f'Geographic expansion: {len([p for p in intelligence.profiles if p.job_postings and any("international" in j.location.lower() if j.location else False for j in p.job_postings)])} competitors show international hiring',
            f'Technology adoption: Average of {sum(len(p.technology_stack) if p.technology_stack else 0 for p in intelligence.profiles) / len(intelligence.profiles):.1f} technologies per competitor'
        ]
        
        for dynamic in dynamics:
            doc.add_paragraph(dynamic, style='List Bullet')
    
    def _create_competitor_profile(self, doc, profile):
        """Create individual competitor profile section"""
        from docx.shared import Inches
        
        doc.add_heading(f'{profile.name}', level=1)
        
        # Basic information table
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Shading Accent 1'

        # Table content
        table.cell(0, 0).text = 'Website'
        table.cell(0, 1).text = profile.website

        table.cell(1, 0).text = 'Threat Level'
        table.cell(1, 1).text = profile.threat_level.value.title() if profile.threat_level else 'Medium'

        table.cell(2, 0).text = 'Target Markets'
        table.cell(2, 1).text = ', '.join(profile.target_markets) if profile.target_markets else 'Unknown'

        table.cell(3, 0).text = 'Last Analyzed'
        last_analyzed_value = profile.last_analyzed
        if isinstance(last_analyzed_value, datetime):
            last_analyzed_text = last_analyzed_value.isoformat()
        else:
            last_analyzed_text = last_analyzed_value or 'N/A'
        table.cell(3, 1).text = last_analyzed_text

        table.cell(4, 0).text = 'Technology Stack'
        table.cell(4, 1).text = ', '.join(profile.technology_stack) if profile.technology_stack else 'Unknown'
        
        # Make first column bold
        for i in range(5):
            for paragraph in table.cell(i, 0).paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Funding information
        if profile.funding_info:
            doc.add_heading('Funding Overview', level=2)
            
            funding_table = doc.add_table(rows=4, cols=2)
            funding_table.style = 'Light List Accent 1'
            
            funding_table.cell(0, 0).text = 'Total Funding'
            funding_table.cell(0, 1).text = profile.funding_info.total_funding or 'Unknown'
            
            funding_table.cell(1, 0).text = 'Last Round Amount'
            funding_table.cell(1, 1).text = profile.funding_info.last_round_amount or 'Unknown'
            
            funding_table.cell(2, 0).text = 'Last Round Type'
            funding_table.cell(2, 1).text = profile.funding_info.last_round_type or 'Unknown'
            
            funding_table.cell(3, 0).text = 'Last Round Date'
            funding_table.cell(3, 1).text = profile.funding_info.last_round_date or 'Unknown'
            
            # Make first column bold
            for i in range(4):
                for paragraph in funding_table.cell(i, 0).paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        
        # Key features
        if profile.key_features:
            doc.add_heading('Key Features & Capabilities', level=2)
            for feature in profile.key_features[:10]:  # Limit to top 10
                doc.add_paragraph(feature, style='List Bullet')
        
        # Recent activity
        if profile.job_postings or profile.recent_news or profile.github_activity:
            doc.add_heading('Recent Activity', level=2)
            activity_items = []
            
            if profile.job_postings:
                activity_items.append(f'{len(profile.job_postings)} active job postings')
            if profile.recent_news:
                activity_items.append(f'{len(profile.recent_news)} recent news mentions')
            if profile.github_activity and profile.github_activity.public_repos:
                activity_items.append(f'{profile.github_activity.public_repos} public repositories')
            if profile.case_studies:
                activity_items.append(f'{len(profile.case_studies)} customer case studies')
            
            for item in activity_items:
                doc.add_paragraph(item, style='List Bullet')
    
    async def _create_recommendations(self, doc, intelligence):
        """Create strategic recommendations section"""
        doc.add_heading('Strategic Recommendations', level=1)
        
        # Generate recommendations
        recommendations = await self._generate_recommendations_content(intelligence)
        
        # Parse and format recommendations
        if 'IMMEDIATE ACTIONS' in recommendations:
            sections = recommendations.split('IMMEDIATE ACTIONS')[1].split('SHORT-TERM')[0] if 'SHORT-TERM' in recommendations else recommendations.split('IMMEDIATE ACTIONS')[1]
            doc.add_heading('Immediate Actions (0-90 days)', level=2)
            for line in sections.split('\n'):
                line = line.strip()
                if line and line.startswith('•'):
                    doc.add_paragraph(line, style='List Bullet')
        
        if 'SHORT-TERM' in recommendations:
            sections = recommendations.split('SHORT-TERM')[1].split('LONG-TERM')[0] if 'LONG-TERM' in recommendations else recommendations.split('SHORT-TERM')[1]
            doc.add_heading('Short-term Initiatives (3-6 months)', level=2)
            for line in sections.split('\n'):
                line = line.strip()
                if line and line.startswith('•'):
                    doc.add_paragraph(line, style='List Bullet')
        
        if 'LONG-TERM' in recommendations:
            sections = recommendations.split('LONG-TERM')[1].split('CONTINUOUS')[0] if 'CONTINUOUS' in recommendations else recommendations.split('LONG-TERM')[1]
            doc.add_heading('Long-term Strategy (6-12 months)', level=2)
            for line in sections.split('\n'):
                line = line.strip()
                if line and line.startswith('•'):
                    doc.add_paragraph(line, style='List Bullet')
    
    def _create_appendix(self, doc, intelligence):
        """Create appendix section"""
        doc.add_heading('Appendix', level=1)
        
        # Data sources used
        doc.add_heading('Data Sources', level=2)
        if intelligence.metadata and 'data_sources_used' in intelligence.metadata:
            sources = intelligence.metadata['data_sources_used']
            for source in sources:
                doc.add_paragraph(f'{source.title()}', style='List Bullet')
        
        # Analysis configuration
        doc.add_heading('Analysis Configuration', level=2)
        if intelligence.config:
            config_table = doc.add_table(rows=5, cols=2)
            config_table.style = 'Light List Accent 1'
            
            config_table.cell(0, 0).text = 'Analysis Depth'
            config_table.cell(0, 1).text = intelligence.config.depth_level.value.title()
            
            config_table.cell(1, 0).text = 'Competitors'
            config_table.cell(1, 1).text = ', '.join(intelligence.config.competitors)
            
            config_table.cell(2, 0).text = 'Output Formats'
            config_table.cell(2, 1).text = ', '.join(intelligence.config.output_formats)
            
            config_table.cell(3, 0).text = 'Analysis Date'
            config_table.cell(3, 1).text = intelligence.analysis_date
            
            config_table.cell(4, 0).text = 'Total Profiles'
            config_table.cell(4, 1).text = str(len(intelligence.profiles))
            
            # Make first column bold
            for i in range(5):
                for paragraph in config_table.cell(i, 0).paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
    
    # Helper methods (reuse from PDF generator)
    def _get_funding_status(self, profile) -> str:
        """Get funding status for threat matrix"""
        if not profile.funding_info or not profile.funding_info.total_funding:
            return 'Unknown'
        
        funding = profile.funding_info.total_funding
        if 'B' in funding:
            return 'Well-funded (Billion+)'
        elif 'M' in funding:
            try:
                amount = float(funding.replace('$', '').replace('M', '').replace(',', ''))
                if amount > 200:
                    return 'Well-funded (200M+)'
                elif amount > 50:
                    return 'Funded (50M+)'
                else:
                    return 'Early-stage'
            except:
                return 'Funded'
        return 'Early-stage'
    
    def _get_key_strengths(self, profile) -> str:
        """Get key strengths for threat matrix"""
        strengths = []
        
        if profile.key_features and len(profile.key_features) > 15:
            strengths.append("Comprehensive features")
        
        if profile.case_studies and len(profile.case_studies) > 8:
            strengths.append("Strong customer base")
        
        if profile.funding_info and profile.funding_info.total_funding and 'B' in profile.funding_info.total_funding:
            strengths.append("Well-funded")
        
        if profile.technology_stack:
            ai_tech = [t for t in profile.technology_stack if t in ['AI', 'Machine Learning']]
            if ai_tech:
                strengths.append("AI capabilities")
        
        return '; '.join(strengths[:2]) if strengths else 'Standard offering'
    
    def _get_primary_concerns(self, profile) -> str:
        """Get primary concerns for threat matrix"""
        concerns = []
        
        if profile.threat_level and profile.threat_level.value in ['high', 'critical']:
            concerns.append("Direct competitive threat")
        
        if profile.job_postings and len(profile.job_postings) > 15:
            concerns.append("Aggressive hiring")
        
        if profile.recent_news and len(profile.recent_news) > 5:
            concerns.append("High market activity")
        
        if profile.funding_info and profile.funding_info.last_round_date:
            from datetime import datetime
            try:
                last_round = datetime.fromisoformat(profile.funding_info.last_round_date.replace('Z', '+00:00'))
                if (datetime.now() - last_round).days < 365:
                    concerns.append("Recent funding")
            except:
                pass
        
        return '; '.join(concerns[:2]) if concerns else 'Moderate activity'
    
    async def _generate_executive_summary_content(self, intelligence) -> str:
        """Generate executive summary using LLM"""
        try:
            competitor_names = [p.name for p in intelligence.profiles]
            high_threat_count = len([p for p in intelligence.profiles 
                                   if p.threat_level and p.threat_level.value in ['high', 'critical']])
            
            prompt = f"""
            Create an executive summary for a competitive analysis of {len(intelligence.profiles)} ecommerce search competitors: {', '.join(competitor_names)}.
            
            Key findings:
            - {high_threat_count} competitors pose high/critical threat
            - Market analysis covers funding, features, and market positioning
            - Analysis includes technology trends and innovation indicators
            
            Structure the summary with:
            1. Market Overview (2-3 sentences)
            2. Primary Competitive Threats (3 bullet points)
            3. Strategic Opportunities (2-3 bullet points)
            4. Immediate Action Items (3 bullet points with timeframes)
            
            Keep under 400 words, executive-friendly, and actionable.
            """
            
            model = self.config.get_llm_model('summary')
            summary = self.llm_provider.chat(
                system="You are a senior business analyst creating executive summaries for C-level executives.",
                user=prompt,
                model=model
            )
            
            return summary
            
        except Exception as e:
            logger.warning(f"LLM summary generation failed: {e}")
            return self._create_fallback_summary(intelligence)
    
    def _create_fallback_summary(self, intelligence) -> str:
        """Create fallback summary without LLM"""
        competitor_count = len(intelligence.profiles)
        high_threat_count = len([p for p in intelligence.profiles 
                               if p.threat_level and p.threat_level.value in ['high', 'critical']])
        
        return f"""
        This analysis examines {competitor_count} key competitors in the ecommerce search market, providing strategic insights for competitive positioning and market response.
        
        Key findings include {high_threat_count} competitors posing high competitive threat requiring immediate attention. The market shows continued consolidation with active funding and feature development, while AI/ML capabilities are becoming standard across leading vendors.
        
        Strategic implications suggest accelerated product development is needed to maintain competitive parity, pricing pressure will likely increase as the market matures, and customer retention becomes critical as competitors strengthen their offerings.
        
        Recommended actions include conducting detailed feature gap analysis within 90 days, evaluating pricing strategy competitiveness by end of quarter, and strengthening customer success and retention programs immediately.
        """
    
    async def _generate_market_analysis_content(self, intelligence) -> str:
        """Generate market analysis content"""
        try:
            market_data = {
                'total_competitors': len(intelligence.profiles),
                'well_funded_count': len([p for p in intelligence.profiles 
                                        if p.funding_info and p.funding_info.total_funding and 'B' in p.funding_info.total_funding]),
                'high_threat_count': len([p for p in intelligence.profiles 
                                        if p.threat_level and p.threat_level.value in ['high', 'critical']]),
                'avg_features': sum(len(p.key_features) if p.key_features else 0 for p in intelligence.profiles) / len(intelligence.profiles)
            }
            
            prompt = f"""
            Analyze the competitive landscape for ecommerce search based on these market indicators:
            
            Market Structure:
            - {market_data['total_competitors']} major competitors analyzed
            - {market_data['well_funded_count']} billion-dollar funded companies
            - {market_data['high_threat_count']} pose high competitive threat
            - Average of {market_data['avg_features']:.1f} features per competitor
            
            Provide analysis covering:
            1. Market maturity and competitive intensity
            2. Key competitive dynamics and trends
            3. Market consolidation vs fragmentation
            4. Innovation trends across competitors
            
            Keep analysis strategic and under 300 words.
            """
            
            model = self.config.get_llm_model('analysis')
            analysis = self.llm_provider.chat(
                system="You are a market research analyst specializing in B2B SaaS competitive landscapes.",
                user=prompt,
                model=model
            )
            
            return analysis
            
        except Exception as e:
            logger.warning(f"Market analysis generation failed: {e}")
            return self._create_fallback_market_analysis(intelligence)
    
    def _create_fallback_market_analysis(self, intelligence) -> str:
        """Create fallback market analysis"""
        competitor_count = len(intelligence.profiles)
        well_funded = len([p for p in intelligence.profiles 
                          if p.funding_info and p.funding_info.total_funding])
        
        return f"""
        The ecommerce search market shows high competitive intensity with {competitor_count} major players analyzed. 
        
        Market maturity is evidenced by the presence of {well_funded} well-funded competitors, indicating established players and significant barriers to entry.
        
        Competitive dynamics are primarily feature-driven, with most vendors offering comprehensive search capabilities. Differentiation focuses on AI/ML capabilities, ease of integration, and vertical specialization.
        
        Innovation trends show artificial intelligence and machine learning becoming table stakes, with advanced personalization and real-time capabilities emerging as key differentiators.
        
        The market is consolidating around well-funded leaders while maintaining room for specialized players with unique value propositions.
        """
    
    async def _generate_recommendations_content(self, intelligence) -> str:
        """Generate strategic recommendations"""
        try:
            high_threat_competitors = [p.name for p in intelligence.profiles 
                                     if p.threat_level and p.threat_level.value in ['high', 'critical']]
            
            prompt = f"""
            Based on competitive analysis of ecommerce search competitors, provide strategic recommendations.
            
            Context:
            - {len(intelligence.profiles)} competitors analyzed
            - High threat competitors: {', '.join(high_threat_competitors) if high_threat_competitors else 'None identified'}
            - Market shows active funding and feature development
            
            Provide 5-7 specific, actionable recommendations with:
            1. Clear action items
            2. Suggested timeframes (90 days, 6 months, 1 year)
            3. Business rationale
            
            Focus on product strategy, competitive positioning, and market response.
            """
            
            model = self.config.get_llm_model('analysis')
            recommendations = self.llm_provider.chat(
                system="You are a strategic business consultant providing competitive response recommendations.",
                user=prompt,
                model=model
            )
            
            return recommendations
            
        except Exception as e:
            logger.warning(f"Recommendations generation failed: {e}")
            return self._create_fallback_recommendations(intelligence)
    
    def _create_fallback_recommendations(self, intelligence) -> str:
        """Create fallback recommendations"""
        return """
        IMMEDIATE ACTIONS (0-90 days):
        • Conduct detailed feature gap analysis against top 3 competitors
        • Strengthen customer retention programs and success metrics
        • Evaluate pricing strategy competitiveness
        
        SHORT-TERM INITIATIVES (3-6 months):
        • Accelerate AI/ML capability development to match market standards
        • Enhance integration ecosystem and developer experience
        • Implement competitive monitoring and intelligence system
        
        LONG-TERM STRATEGY (6-12 months):
        • Consider strategic partnerships or acquisitions for capability gaps
        • Develop vertical-specific solutions for market differentiation
        • Invest in next-generation search technologies (voice, visual, semantic)
        
        CONTINUOUS ACTIVITIES:
        • Monitor competitor funding rounds and strategic moves
        • Track feature releases and market positioning changes
        • Maintain competitive battle cards for sales team
        """
    
    async def _generate_simple_text_report(self, intelligence) -> str:
        """Fallback to text report if DOCX unavailable"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"competitor_analysis_{timestamp}.txt"
        filepath = self.output_dir / filename
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write("COMPETITOR ANALYSIS REPORT\n")
            f.write("=" * 50 + "\n\n")
            f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"Competitors Analyzed: {len(intelligence.profiles)}\n\n")
            
            # Executive Summary
            f.write("EXECUTIVE SUMMARY\n")
            f.write("-" * 20 + "\n")
            summary = await self._generate_executive_summary_content(intelligence)
            f.write(summary + "\n\n")
            
            # Competitor Profiles
            f.write("COMPETITOR PROFILES\n")
            f.write("-" * 20 + "\n")
            for profile in intelligence.profiles:
                f.write(f"\n{profile.name}\n")
                f.write("=" * len(profile.name) + "\n")
                f.write(f"Website: {profile.website}\n")
                f.write(f"Threat Level: {profile.threat_level.value if profile.threat_level else 'Medium'}\n")
                
                if profile.funding_info and profile.funding_info.total_funding:
                    f.write(f"Funding: {profile.funding_info.total_funding}\n")
                
                if profile.key_features:
                    f.write(f"Key Features: {len(profile.key_features)} features identified\n")
                
                if profile.job_postings:
                    f.write(f"Active Hiring: {len(profile.job_postings)} job postings\n")
                
                f.write("\n")
            
            # Recommendations
            f.write("STRATEGIC RECOMMENDATIONS\n")
            f.write("-" * 25 + "\n")
            recommendations = await self._generate_recommendations_content(intelligence)
            f.write(recommendations + "\n")
        
        logger.info(f"Text report generated: {filepath}")
        return str(filepath)